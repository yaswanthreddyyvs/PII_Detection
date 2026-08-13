from pathlib import Path
from typing import Callable

from docx import Document


ReplacementFunction = Callable[[str], str]


class DocumentWriter:
    """Writes redacted text into a DOCX document."""

    def write(
        self,
        input_path: Path,
        output_path: Path,
        redact_text: ReplacementFunction,
    ) -> None:

        document = Document(input_path)

        for paragraph in document.paragraphs:
            paragraph.text = redact_text(paragraph.text)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.text = redact_text(paragraph.text)

        for section in document.sections:
            for paragraph in section.header.paragraphs:
                paragraph.text = redact_text(paragraph.text)

            for paragraph in section.footer.paragraphs:
                paragraph.text = redact_text(paragraph.text)

        output_path.parent.mkdir(
            parents=True,
            exist_ok=True,
        )

        document.save(output_path)